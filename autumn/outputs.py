

# external imports
import openpyxl as xl
from docx import Document
from matplotlib import pyplot, patches
from matplotlib.ticker import FuncFormatter
import numpy
import platform
import os
import economics
import copy
from scipy import stats
import itertools
import shutil

# AuTuMN import
import tool_kit as t_k


''' plot creating and cleaning functions '''


def initialise_figures_axes(n_panels, room_for_legend=False, requested_grid=None, share_yaxis='none'):
    """
    Initialise the subplots (or single plot) according to the number of panels required.

    Args:
        n_panels: The number of panels needed
        room_for_legend: Whether room is needed for a legend - applies to single axis plots only
        requested_grid: Shape of grid panels requested at call to method
        share_yaxis: String to pass to the sharey option
    Returns:
        fig: The figure object
        axes: A list containing each of the axes
        max_dims: The number of rows or columns of sub-plots, whichever is greater
    """

    pyplot.style.use('ggplot')
    n_rows, n_cols = requested_grid if requested_grid else find_subplot_grid(n_panels)
    horizontal_position_one_axis = .08 if room_for_legend else .15
    if n_panels == 1:
        fig = pyplot.figure()
        axes = fig.add_axes([horizontal_position_one_axis, .15, 0.7, 0.7])
    elif n_panels == 2:
        fig, axes = pyplot.subplots(1, 2)
        fig.set_figheight(3.5)
        fig.subplots_adjust(bottom=.15, top=.85)
    else:
        fig, axes = pyplot.subplots(n_rows, n_cols, sharey=share_yaxis)
        for panel in range(n_panels, n_rows * n_cols):
            find_panel_grid_indices(axes, panel, n_rows, n_cols).axis('off')
    return fig, axes, max([n_rows, n_cols]), n_rows, n_cols


def add_title_to_plot(fig, n_panels, content):
    """
    Function to add title to the top of a figure and handle multiple panels if necessary.

    Args:
        fig: The figure object to have a title added to it
        n_panels: Integer for the total number of panels on the figure
        content: Unprocessed string to determine text for the title
    """

    # if few panels, bigger and lower title
    greater_heights = {1: .92, 2: .98}
    greater_font_sizes = {1: 14, 2: 11}
    fig.suptitle(t_k.find_title_from_dictionary(content),
                 y=greater_heights[n_panels] if n_panels in greater_heights else .96,
                 fontsize=greater_font_sizes[n_panels] if n_panels in greater_font_sizes else 10)


def get_label_font_size(max_dim):
    """
    Find standardised font size that can be applied across all figures.

    Args:
        max_dim: The number of rows or columns, whichever is the greater
    """

    label_font_sizes = {1: 8, 2: 7}
    return label_font_sizes[max_dim] if max_dim in label_font_sizes else 6


def scale_axes(vals, max_val, y_sig_figs):
    """
    General function to scale a set of axes and produce text that can be added to the axis label. Written here as a
    separate function from the tidy_axis method below because it can then be applied to both x- and y-axes.

    Args:
        vals: List of the current y-ticks
        max_val: The maximum value of this list
        y_sig_figs: The preferred number of significant figures for the ticks
    Returns:
        labels: List of the modified tick labels
        axis_modifier: The text to be added to the axis
    """

    y_number_format = '%.' + str(y_sig_figs) + 'f'
    y_number_format_around_one = '%.' + str(max(2, y_sig_figs)) + 'f'
    if max_val < 5e-9:
        labels = [y_number_format % (v * 1e12) for v in vals]
        axis_modifier = 'Trillionth '
    elif max_val < 5e-6:
        labels = [y_number_format % (v * 1e9) for v in vals]
        axis_modifier = 'Billionth '
    elif max_val < 5e-3:
        labels = [y_number_format % (v * 1e6) for v in vals]
        axis_modifier = 'Millionth '
    elif max_val < 5e-2:
        labels = [y_number_format % (v * 1e3) for v in vals]
        axis_modifier = 'Thousandth '
    elif max_val < .1:
        labels = [y_number_format % (v * 1e2) for v in vals]
        axis_modifier = 'Hundredth '
    elif max_val < 5:
        labels = [y_number_format_around_one % v for v in vals]
        axis_modifier = ''
    elif max_val < 5e3:
        labels = [y_number_format % v for v in vals]
        axis_modifier = ''
    elif max_val < 5e6:
        labels = [y_number_format % (v / 1e3) for v in vals]
        axis_modifier = 'Thousand '
    elif max_val < 5e9:
        labels = [y_number_format % (v / 1e6) for v in vals]
        axis_modifier = 'Million '
    else:
        labels = [y_number_format % (v / 1e9) for v in vals]
        axis_modifier = 'Billion '
    return labels, axis_modifier


def add_legend_to_plot(axis, max_dim, location=0):
    """
    Add legend to plot, with font size determined by the maximum number of dimensions of subplot panels.

    Args:
        axis: The axis to have the legend added
        max_dim: The number of rows or columns of subplots, whichever is the greater
        location: The matplotlib integer specifying the position for the legend (default of zero is 'best')
    """

    if max_dim == 1:
        axis.legend(bbox_to_anchor=(1.3, 1), fontsize=get_label_font_size(max_dim))
    else:
        axis.legend(fontsize=get_label_font_size(max_dim), loc=location)


''' plot navigation functions '''


def find_subplot_grid(n_plots):
    """
    Find a convenient number of rows and columns for a required number of subplots. First take the root of the number of
    subplots and round up to find the smallest square that could accommodate all of them. Next find out how many rows
    that many subplots would fill out by dividing the number of plots by the number of columns and rounding up. This
    will potentially leave a few panels blank at the end and number of rows will equal the number of columns or the
    number of rows will be on fewer.

    Args:
        n_plots: The number of subplots needed
    Returns:
        The number of rows of subplots
        n_cols: The number of columns of subplots
    """

    n_cols = int(numpy.ceil(numpy.sqrt(n_plots)))
    return int(numpy.ceil(n_plots / float(n_cols))), n_cols


def find_panel_grid_indices(axes, index, n_rows, n_columns):
    """
    Find the subplot index for a plot panel from the number of the panel and the number of columns of sub-plots.

    Args:
        axes: All the plot axes to be searched from
        index: The number of the panel counting up from zero
        n_rows: Number of rows of sub-plots in figure
        n_columns: Number of columns of sub-plots in figure
    """

    row, column = numpy.floor_divide(index, n_columns), (index + 1) % n_columns - 1 if n_rows > 1 else None
    return axes[row, column] if n_rows > 1 else axes[index]


def last_row(index, n_rows, n_columns):
    """
    Determine whether panel is not in the last row of sub-plots of the figure from the index of the panel and the number
    of rows and columns of subplots.

    Args:
        index: The number of the panel counting up from zero
        n_rows: Number of rows of sub-plots in the figure
        n_columns: Number of columns of sub-plots in the figure
    """

    return index >= (n_rows - 1) * n_columns


''' patch-related functions '''


def create_patch_from_list(x_values, lower_border, upper_border):
    """
    Creates an array that can be used to plot a patch using the add patch plotting function in matplotlib.

    Args:
        x_values: The x-values to go forward and backward with
        lower_border: The lower edge of the patch
        upper_border: The upper edge of the patch
    Returns:
        patch_array: An array for use in plotting patches
            (with length of double the length of the inputs lists and height of two)
    """

    assert len(x_values) == len(lower_border) == len(upper_border), 'Tried to make patch out of unequal length lists'
    patch_array = numpy.zeros(shape=(len(x_values) * 2, 2))
    for x_num, x_value in enumerate(x_values):
        patch_array[x_num][0] = x_value  # x_values forwards
        patch_array[-x_num - 1][0] = x_value  # x_values backwards
        patch_array[x_num][1] = lower_border[x_num]  # lower limit data forwards
        patch_array[-x_num - 1][1] = upper_border[x_num]  # upper limit data backwards
    return patch_array


def increment_list_for_patch(new_data, cumulative_data):
    """
    Takes a list of cumulative data totals, preserves the previous values and adds a new list to it. This is to allow
    patches to be plotted that have the previous data values as their base and the results of this stacking as their
    top.

    Args:
        new_data: The new data to be stacked up
        cumulative_data: The previous running totals
    Returns:
        previous_data: The previous running total (was cumulative_data)
        The new running total as the new values for cumulative_data
    """

    previous_data = copy.copy(cumulative_data)
    return previous_data, [last + current for last, current in zip(cumulative_data, new_data)]


''' target plotting functions '''


def find_exponential_constants(x_values, y_values):
    """
    Find the parameters to an exponential function that passes through the point (x_values[0], y_values[0])
    and (x_values[1], y_values[1]) and is of the form: y = exp(-a * (x - b)), where x is the independent variable.

    Args:
        x_values: List of the two time or x coordinates of the points to be fitted to
        y_values: List of the two outputs or y coordinates of the points to be fitted to
    Returns:
        Parameter for the horizontal transformation of the function
        b: Parameter for the horizontal translation of the function
    """

    b = (x_values[0] * numpy.log(y_values[1]) - x_values[1] * numpy.log(y_values[0])) \
        / (numpy.log(y_values[1]) - numpy.log(y_values[0]))
    return - numpy.log(y_values[0]) / (x_values[0] - b), b


def plot_endtb_targets(ax, output, base_value, plot_colour):
    """
    Plot the End TB Targets and the direction that we need to head to achieve them.

    Args:
        ax: The axis to be plotted to
        output: Output string
        base_value: The value of the output at the reference time
        plot_colour: List of colours for plotting
    """

    # End TB Targets data to use
    times = [2015., 2020., 2025., 2030., 2035.]
    target_text = ['', 'M', 'M', 'S', 'E']  # M for milestone, S for Sustainable Development Goal, E for End TB Target
    target_props_dict \
        = {'mortality': [1., .65, .25, .1, .05],
           'incidence': [1., .8, .5, .2, .1]}

    # find the actual target values
    target_values = [base_value * target for target in target_props_dict[output]]

    # plot the individual targets themselves
    ax.plot(times[1:], target_values[1:], marker='o', markersize=4, color=plot_colour, markeredgewidth=0., linewidth=0.)

    # cycle through times and plot
    for t in range(len(times) - 1):
        times_to_plot, output_to_reach_target \
            = find_values_exp_function([times[t], times[t + 1]], [target_values[t], target_values[t + 1]])
        ax.plot(times_to_plot, output_to_reach_target, color=plot_colour, linewidth=.5)

    # annotate points with letters if requested
    for t, text in enumerate(target_props_dict[output]):
        ax.annotate(target_text[t], (times[t], target_values[t] + target_values[1] / 20.),
                    horizontalalignment='center', verticalalignment='bottom', fontsize=8, color=plot_colour)


def find_values_exp_function(times, target_values, number_x_values=100):
    """
    Find the times to plot and the outputs tracking towards the targets from the list of times and target values,
    using the function to fit exponential functions.

    Args:
        times: The list of times being worked through
        target_values: The list fo target values corresponding to the times
        number_x_values: The number of x-values required for plotting
    Returns:
        times_to_plot: List of the x-values or times for plotting
        Corresponding list of values for the output needed to track towards the target
    """

    a, b = find_exponential_constants([times[0], times[1]], [target_values[0], target_values[1]])
    times_to_plot = numpy.linspace(times[0], times[1], number_x_values)
    return times_to_plot, [numpy.exp(-a * (x - b)) for x in times_to_plot]


''' spreadsheet writing functions '''


def write_param_to_sheet(country_sheet, working_list, median_run_index):
    """
    Function to write a single parameter value into a cell of the input spreadsheets - to be used as part of
    automatic calibration.

    Args:
        country_sheet: Spreadsheet object
        working_list: List to take value from
        median_run_index: Integer index of the median
    """

    for param in working_list:
        if working_list[param]:

            # find value to write from list and index
            value = working_list[param][median_run_index]

            # over-write existing parameter value if present
            param_found = False
            for row in country_sheet.rows:
                if row[0].value == param:
                    row[1].value, param_found = value, True

            # if parameter not found in existing spreadsheet, write into new row at the bottom
            if not param_found:
                max_row = country_sheet.max_row
                country_sheet.cell(row=max_row + 1, column=1).value = param
                country_sheet.cell(row=max_row + 1, column=2).value = value


def reverse_inputs_if_required(inputs, condition):
    """
    Very simple function to reverse a list if requested, but used so frequently during spreadsheet writing that worth
    having.

    Args:
        inputs: A list of the inputs
        condition: Boolean for whether to reverse or not
    Returns:
        The list reversed if condition and the list unchanged otherwise
    """

    return inputs[::-1] if condition else inputs


class Project:
    def __init__(self, runner, gui_inputs, out_dir_project=None):
        """
        Initialises an object of class Project, that will contain all the information (data + outputs) for writing a
        report for a country.

        Args:
            runner: The main model runner object used to execute all the analyses
            gui_inputs: All inputs from the graphical user interface
        """

        self.model_runner = runner
        self.gui_inputs = gui_inputs
        self.country = self.gui_inputs['country'].lower()
        self.out_dir_project = os.path.join('projects', 'test_' + self.country) \
            if not out_dir_project else out_dir_project
        if not os.path.isdir(self.out_dir_project):
            os.makedirs(self.out_dir_project)
        (self.inputs, self.run_mode) \
            = [None for _ in range(2)]
        (self.output_colours, self.uncertainty_output_colours, self.classified_scaleups,
         self.outputs, self.interpolated_uncertainty) \
            = [{} for _ in range(5)]
        (self.grid, self.plot_rejected_runs, self.plot_true_outcomes) \
            = [False for _ in range(3)]
        (self.accepted_no_burn_in_indices, self.scenarios, self.interventions_to_cost, self.accepted_indices,
         self.accepted_run_weights) = [[] for _ in range(5)]
        self.uncertainty_centiles = {'epi': {}, 'cost': {}}
        for attribute in ['inputs', 'outputs']:
            setattr(self, attribute, getattr(self.model_runner, attribute))
        for attribute in ['scenarios', 'interventions_to_cost', 'run_mode']:
            setattr(self, attribute, getattr(self.model_runner.inputs, attribute))
        self.figure_number, self.title_size = 1, 13
        self.figure_formats = ['png']   # list of strings to allow for multiple formats. e.g. ['png', 'pdf']
        self.years_to_write \
            = range(int(self.inputs.model_constants['report_start_time']),
                    int(self.inputs.model_constants['report_end_time']),
                    int(self.inputs.model_constants['report_step_time']))
        self.classifications \
            = ['demo_', 'econ_', 'epi_prop_smear', 'program_prop_', 'program_timeperiod_',
               'program_prop_novel', 'program_prop_treatment', 'program_prop_detect',
               'int_prop_vaccination', 'program_prop_treatment_success',
               'program_prop_treatment_death', 'transmission_modifier', 'algorithm']
        self.quantities_to_write_back = ['all_parameters', 'all_compartment_values', 'adjustments']
        self.gtb_available_outputs = ['incidence', 'mortality', 'prevalence', 'notifications']
        self.level_conversion_dict = {'lower_limit': '_lo', 'upper_limit': '_hi', 'point_estimate': ''}

        # common times to interpolate uncertainty to
        self.n_interpolation_points = int(2000)
        start_interpolation_time = self.inputs.model_constants['early_time'] if self.run_mode == 'int_uncertainty' \
            else self.inputs.model_constants['early_time']

        self.interpolation_times_uncertainty \
            = numpy.linspace(start_interpolation_time, self.inputs.model_constants['report_end_time'],
                             self.n_interpolation_points)

        # comes up so often that we need to find this index, that easiest to do in instantiation
        self.manual_baseline_start_index = self.find_start_time_index(self.inputs.model_constants['plot_start_time'], 0)

        # standard graphing themes
        self.tick_length = 3
        self.colour_theme \
            = [(0., 0., 0.),
               (0., 0., 125. / 255.),
               (210. / 255., 70. / 255., 0.),
               (100. / 255., 150. / 255., 1.),
               (65. / 255., 65. / 255., 65. / 255.),
               (220. / 255., 25. / 255., 25. / 255.),
               (120. / 255., 55. / 255., 20. / 255.),
               (120. / 255., 55. / 255., 110. / 255.),
               (135. / 255., 135. / 255., 30. / 255.),
               (120. / 255., 120. / 255., 120. / 255.),
               (220. / 255., 20. / 255., 170. / 255.),
               (20. / 255., 65. / 255., 20. / 255.),
               (15. / 255., 145. / 255., 25. / 255.),
               (15. / 255., 185. / 255., 240. / 255.),
               (10. / 255., 0., 110. / 255.)]
        self.gtb_indices \
            = {'incidence': 'e_inc_100k',
               'mortality': 'e_mort_exc_tbhiv_100k',
               'prevalence': 'e_prev_100k'}
        self.gtb_patch_colours \
            = {'incidence': self.colour_theme[1],
               'mortality': self.colour_theme[2],
               'prevalence': self.colour_theme[3],
               'notifications': (0., 0., 0.)}

    ''' master method to call the others '''

    def master_outputs_runner(self):
        """
        Method to work through all the fundamental output methods, which then call all the specific output methods for
        plotting and writing as required.
        """

        self.model_runner.add_comment_to_gui_window('Creating outputs')

        # processing methods that are only required for outputs
        if self.run_mode == 'epi_uncertainty':
            self.find_uncertainty_indices()
            self.calculate_accepted_run_weights()
            self.uncertainty_centiles['epi'] = self.find_uncertainty_common_times('epi')
        elif self.run_mode == 'int_uncertainty':
            for output_type in ['epi', 'cost']:
                self.uncertainty_centiles[output_type] = self.find_uncertainty_common_times(output_type)

        # write automatic calibration values back to sheets
        if self.run_mode == 'epi_uncertainty' and self.gui_inputs['write_uncertainty_outcome_params']:
            self.write_automatic_calibration_outputs()

        # write spreadsheets with sheet for each scenario or each output
        if self.gui_inputs['output_spreadsheets']:
            self.model_runner.add_comment_to_gui_window('Writing output spreadsheets')
            if self.gui_inputs['output_by_scenario']:
                self.write_xls_by_scenario()
            else:
                self.write_xls_by_output()

        # write documents - with document for each scenario or each output
        if self.gui_inputs['output_documents']:
            self.model_runner.add_comment_to_gui_window('Writing output documents')
            if self.gui_inputs['output_by_scenario']:
                self.write_docs_by_scenario()
                if self.run_mode == 'int_uncertainty':
                    self.print_int_uncertainty_relative_change(year=2035.)
            else:
                self.write_docs_by_output()

        # master plotting method
        self.model_runner.add_comment_to_gui_window('Creating plot figures')
        self.run_plotting()

        return os.path.abspath(self.out_dir_project)

    ''' general methods for use by specific methods below '''

    def find_start_time_index(self, time, scenario, by_run=False, run=0, purpose=None):
        """
        There are various messy legacy approaches to this, but now trying to reconcile all these approaches to finding
        the starting index for plotting into one method.
        """

        mode = purpose if purpose else self.run_mode
        if (mode == 'scenario' and scenario) or mode == 'int_uncertainty':
            return 0
        elif mode == 'epi_uncertainty' and by_run:
            times_to_search = self.outputs[mode]['epi'][scenario]['times'][run]
        elif mode == 'epi_uncertainty':
            times_to_search = self.interpolation_times_uncertainty
        else:
            times_to_search = self.outputs['manual']['epi'][scenario]['times']
        return t_k.find_first_list_element_at_least(times_to_search, time)

    def tidy_x_axis(self, axis, start, end, max_dim, labels_off=False, x_label=None):
        """
        Function to tidy x-axis of a plot panel - currently only used in the scale-up vars, but intended to be written
        in such a way as to be extendable to other types of plotting.

        Args:
            axis: The plotting axis
            start: Lowest x-value being plotted
            end: Highest x-value being plotted
            max_dim: Maximum number of rows or columns of subplots in figure
            labels_off: Whether to turn all tick labels off on this axis
            x_label: Text for the x-axis label if required
        """

        # range
        axis.set_xlim(left=start, right=end)

        # ticks and their labels
        if labels_off:
            axis.tick_params(axis='x', labelbottom='off')
        elif len(axis.get_xticks()) > 7:
            for label in axis.xaxis.get_ticklabels()[::2]:
                label.set_visible(False)
        axis.tick_params(axis='x', length=self.tick_length, pad=6, labelsize=get_label_font_size(max_dim))

        # axis label
        if x_label:
            axis.set_xlabel(x_label, fontsize=get_label_font_size(max_dim))

    def tidy_y_axis(self, axis, quantity, max_dims, left_axis=True, max_value=1e6, space_at_top=.1, y_label=None,
                    y_lims=None, allow_negative=False):
        """
        General approach to tidying up the vertical axis of a plot, depends on whether it is the left-most panel.

        Args:
            axis: The axis itself
            quantity: The name of the quantity being plotted (which can be used to determine the sort of variable it is)
            max_dims: Maximum number of rows or columns of subplots on the figure
            left_axis: Boolean for whether the axis is the left-most panel
            max_value: The maximum value in the data being plotted
            space_at_top: Relative amount of space to leave at the top, above the maximum value of the plotted data
            y_label: A label for the y-axis, if required
            y_lims: 2-element tuple for the y-limit, if required
            allow_negative: Whether to set the bottom of the axis to zero
        """

        # axis range
        if y_lims:
            axis.set_ylim(y_lims)
        elif 'prop_' in quantity and axis.get_ylim()[1] > 1.:
            axis.set_ylim(top=1.004)
        elif 'prop_' in quantity or 'likelihood' in quantity or 'cost' in quantity:
            pass
        elif axis.get_ylim()[1] < max_value * (1. + space_at_top):
            axis.set_ylim(top=max_value * (1. + space_at_top))
        if not allow_negative:
            axis.set_ylim(bottom=0.)

        # ticks
        axis.tick_params(axis='y', length=self.tick_length, pad=6, labelsize=get_label_font_size(max_dims))

        # tick labels
        if not left_axis:
            pyplot.setp(axis.get_yticklabels(), visible=False)
        elif 'prop_' in quantity:
            axis.yaxis.set_major_formatter(FuncFormatter('{0:.0%}'.format))

        # axis label
        if y_label and left_axis:
            axis.set_ylabel(y_label, fontsize=get_label_font_size(max_dims))

    def finish_off_figure(self, fig, n_plots, end_filename, title_text):
        """
        Slight extension of save_figure to include adding main title to figure.

        Args:
            fig: The figure to add the title to
            n_plots: The total number of plots in the figure
            end_filename: The end of the string for the file name
            title_text: Text for the title of the figure
        """

        if self.gui_inputs['plot_option_title']:
            add_title_to_plot(fig, n_plots, title_text)
        for file_format in self.figure_formats:
            filename = os.path.join(self.out_dir_project, self.country + end_filename + '.' + file_format)
            fig.savefig(filename, dpi=300)

    def find_uncertainty_indices(self):
        """
        Quick method to create a list of the indices of interest for the runs of the uncertainty analysis.

        Updates:
            self.accepted_no_burn_in_indices: List of the uncertainty indices of interest
        """
        self.accepted_indices = self.outputs['epi_uncertainty']['accepted_indices']
        self.accepted_no_burn_in_indices = [i for i in self.accepted_indices if i >= self.gui_inputs['burn_in_runs']]

    def calculate_accepted_run_weights(self):
        """
        This method accounts for the number of rejections following each acceptance and weights the relevant accepted
        parameter sets accordingly. It will populate the attribute "accepted_run_weights" of the output object.
        """
        for i, accepted_index in enumerate(self.accepted_no_burn_in_indices):
            next_accepted_index = self.accepted_no_burn_in_indices[i+1] \
                if i < (len(self.accepted_no_burn_in_indices) - 1) else accepted_index + 1
            self.accepted_run_weights.append(next_accepted_index - accepted_index)

    def find_uncertainty_common_times(self, output_type):
        """
        Use simple linear interpolation to find the values of outputs from each model run at a standardised set of times
        and then calculate percentiles. Not sure whether this will work for cost outputs - should in theory, but the
        point is mostly about epi outputs anyway.

        Args:
            output_type: Whether epi or cost outputs are being considered
        """

        uncertainty_centiles = {}
        for scenario in self.outputs[self.run_mode][output_type]:
            self.interpolated_uncertainty[scenario] = {}
            uncertainty_centiles[scenario] = {}
            for output in self.outputs[self.run_mode][output_type][scenario]:
                if output != 'times':
                    self.interpolated_uncertainty[scenario][output] \
                        = numpy.empty(shape=(0, self.n_interpolation_points))
                    run_range = range(len(self.outputs['epi_uncertainty']['whether_accepted'])) if \
                        self.run_mode == 'epi_uncertainty' else range(self.inputs.n_samples)
                    for run in run_range:
                        self.interpolated_uncertainty[scenario][output] \
                            = numpy.vstack(
                            (self.interpolated_uncertainty[scenario][output],
                             numpy.interp(self.interpolation_times_uncertainty,
                                          self.outputs[self.run_mode][output_type][scenario]['times'][run, :],
                                          self.outputs[self.run_mode][output_type][scenario][output][run, :])
                             [None, :]))

                    # all runs for scenario analysis (as only accepted recorded) but select accepted ones for baseline
                    matrix_to_analyse = self.interpolated_uncertainty[scenario][output] if scenario \
                        else self.interpolated_uncertainty[scenario][output][self.accepted_no_burn_in_indices]

                    # transform matrix_to_analyse to account for the weights of the accepted runs for epi_uncertainty
                    if self.run_mode == 'epi_uncertainty':
                        matrix_to_analyse = t_k.apply_weighting(matrix_to_analyse, self.accepted_run_weights)

                    # find centiles
                    uncertainty_centiles[scenario][output] \
                        = numpy.percentile(matrix_to_analyse, self.model_runner.percentiles, axis=0)
        return uncertainty_centiles

    def find_uncertainty_centiles(self, mode, output_type):
        """
        Find percentiles from uncertainty dictionaries.

        Args:
            mode: The run mode being considered
            output_type: Whether the output to be calculated is 'epi' or 'cost'
        Updates:
            self.percentiles: Adds all the required percentiles to this dictionary.
        """

        uncertainty_centiles = {}
        for scenario in self.outputs[mode][output_type]:
            uncertainty_centiles[scenario] = {}
            for output in self.outputs[mode][output_type][scenario]:
                if output != 'times':

                    # all runs for scenario analysis (as only accepted recorded) but select accepted ones for baseline
                    matrix_to_analyse = self.outputs[mode][output_type][scenario][output] if scenario \
                        else self.outputs[mode][output_type][scenario][output][self.accepted_no_burn_in_indices]

                    uncertainty_centiles[scenario][output] \
                        = numpy.percentile(matrix_to_analyse, self.model_runner.percentiles, axis=0)
        return uncertainty_centiles

    def sum_compartments_by_category(self, category_type, scenario, start_time_index, fraction=False,
                                     requirements=('',), exclusions=('we all love futsal',), remainder=False):
        """
        Find the sum of all compartments within a particular category (i.e. containing a particular string).

        Args:
            category_type: String for the category of interest
            scenario: Scenario number
            start_time_index: Time index to start from
            fraction: Boolean for whether to calculate as a fraction or as total compartment values
            requirements: List of strings that must appear in the compartments of interest
            exclusions: List of strings of compartments to be ignored
            remainder: Whether to add an additional category for the remainder
        """

        categories_to_loop = getattr(self.inputs, category_type)
        current_data, compartments_in_category = {}, {}
        if remainder:
            current_data['remainder'] = [0.] * len(self.model_runner.models[scenario].times[start_time_index:])
        for label in categories_to_loop:

            # find list of relevant compartments
            compartments_in_category[label] = []
            for comp in self.model_runner.models[scenario].labels:
                if label in comp and all(strings in comp for strings in requirements) \
                        and not any(strings in comp for strings in exclusions):
                    compartments_in_category[label] += [comp]
                elif remainder:
                    compartments_in_category['remainder'] += [comp]

            # sum the values
            current_data[label] = [0.] * len(self.model_runner.models[scenario].times[start_time_index:])
            for compartment in compartments_in_category[label]:
                current_data[label] \
                    = t_k.elementwise_list_addition(current_data[label], self.model_runner.models[
                                                             scenario].compartment_soln[compartment][start_time_index:])
            current_data[label] = t_k.elementwise_list_division(current_data[label], self.outputs['manual']['epi'][
                                        scenario]['population'][start_time_index:]) if fraction else current_data[label]
        return current_data

    ''' methods for outputting to documents and spreadsheets and console '''

    def write_automatic_calibration_outputs(self):
        """
        Write values from automatic calibration process back to input spreadsheets, using the parameter values
        that were associated with the model run with the greatest likelihood.
        """

        try:
            path = os.path.join('autumn/xls/data_' + self.country + '.xlsx')
        except IOError:
            self.model_runner.add_comment_to_gui_window(
                'No country input spreadsheet available for requested uncertainty parameter writing')
        else:
            self.model_runner.add_comment_to_gui_window(
                'Writing automatic calibration parameters back to input spreadsheet')

            # open workbook and sheet
            country_input_book = xl.load_workbook(path)
            country_sheet = country_input_book['constants']

            # find the integration run with the highest likelihood
            best_likelihood_index = self.outputs['epi_uncertainty']['loglikelihoods'].index(
                max(self.outputs['epi_uncertainty']['loglikelihoods']))

            # write the parameters and starting compartment sizes back in to input sheets
            for attribute in self.quantities_to_write_back:
                write_param_to_sheet(country_sheet, self.outputs['epi_uncertainty'][attribute], best_likelihood_index)

            # save
            country_input_book.save(path)

    def write_xls_by_scenario(self):
        """
        Write a spreadsheet with the sheet referring to one scenario.
        """

        # general prelims to work out what to write
        horizontal = self.gui_inputs['output_horizontally']
        result_types = ['epi_', 'raw_cost_', 'inflated_cost_', 'discounted_cost_', 'discounted_inflated_cost_']
        scenarios = [15] if self.run_mode == 'int_uncertainty' else self.scenarios

        # write a new file for each scenario and for each broad category of output
        for result_type in result_types:
            for scenario in scenarios:

                # prepare sheet
                scenario_name = t_k.find_scenario_string_from_number(scenario)
                path = os.path.join(self.out_dir_project, result_type + scenario_name) + '.xlsx'
                workbook = xl.Workbook()
                sheet = workbook.active
                sheet.title = scenario_name
                sheet.cell(row=1, column=1).value = 'Year'

                # year column
                for y, year in enumerate(self.years_to_write):
                    row, column = reverse_inputs_if_required([y + 2, 1], horizontal)
                    sheet.cell(row=row, column=column).value = year

                # epi outputs
                if result_type == 'epi_':

                    # loop over outputs
                    for out, output in enumerate(self.model_runner.epi_outputs_to_analyse):

                        # with uncertainty
                        if self.run_mode == 'epi_uncertainty' or self.run_mode == 'int_uncertainty':

                            # scenario names and confidence interval titles
                            strings_to_write = [t_k.capitalise_and_remove_underscore(output), 'Lower', 'Upper']

                            for ci in range(len(strings_to_write)):
                                row, column = reverse_inputs_if_required([1, out * 3 + 2 + ci], horizontal)
                                sheet.cell(row=row, column=column).value = strings_to_write[ci]

                            # data columns
                            for y, year in enumerate(self.years_to_write):
                                for o in range(3):
                                    row, column = reverse_inputs_if_required([y + 2, out * 3 + 2 + o], horizontal)
                                    sheet.cell(row=row, column=column).value \
                                        = self.uncertainty_centiles['epi'][scenario][output][
                                        o, t_k.find_first_list_element_at_least(self.interpolation_times_uncertainty,
                                                                                year)]

                        # without uncertainty
                        else:

                            # names across top
                            row, column = reverse_inputs_if_required([1, out + 2], horizontal)
                            sheet.cell(row=row, column=column).value = t_k.capitalise_and_remove_underscore(output)

                            # columns of data
                            for y, year in enumerate(self.years_to_write):
                                row, column = reverse_inputs_if_required([y + 2, out + 2], horizontal)
                                sheet.cell(row=row, column=column).value \
                                    = self.outputs['manual']['epi'][scenario][output][
                                        t_k.find_first_list_element_at_least(
                                            self.model_runner.outputs['manual']['epi'][scenario]['times'], year)]

                # economic outputs (uncertainty unavailable)
                elif 'cost_' in result_type:

                    # loop over interventions
                    for inter, intervention in enumerate(self.inputs.interventions_to_cost[scenario]):

                        # names across top
                        row, column = reverse_inputs_if_required([1, inter + 2], horizontal)
                        sheet.cell(row=row, column=column).value = t_k.capitalise_and_remove_underscore(intervention)

                        # data columns
                        for y, year in enumerate(self.years_to_write):
                            row, column = reverse_inputs_if_required([y + 2, inter + 2], horizontal)
                            sheet.cell(row=row, column=column).value \
                                = self.outputs['manual']['cost'][scenario][result_type + intervention][
                                        t_k.find_first_list_element_at_least(
                                            self.model_runner.outputs['manual']['cost'][scenario]['times'], year)]
                workbook.save(path)

    def write_xls_by_output(self):
        """
        Write a spreadsheet with the sheet referring to one output.
        """

        # general prelims to work out what to write
        horizontal = self.gui_inputs['output_horizontally']
        scenarios = [15] if self.run_mode == 'int_uncertainty' else self.scenarios

        # write a new file for each output
        for inter in self.model_runner.epi_outputs_to_analyse:

            # prepare sheet
            path = os.path.join(self.out_dir_project, 'epi_' + inter) + '.xlsx'
            workbook = xl.Workbook()
            sheet = workbook.active
            sheet.title = inter
            sheet.cell(row=1, column=1).value = 'Year'

            # write the year column
            for y, year in enumerate(self.years_to_write):
                row, column = reverse_inputs_if_required([y + 2, 1], horizontal)
                sheet.cell(row=row, column=column).value = year

            # cycle over scenarios
            for s, scenario in enumerate(scenarios):
                scenario_name = t_k.find_scenario_string_from_number(scenario)

                # with uncertainty
                if self.run_mode == 'epi_uncertainty' or self.run_mode == 'int_uncertainty':

                    # scenario names and confidence interval titles
                    strings_to_write = [t_k.capitalise_and_remove_underscore(scenario_name), 'Lower', 'Upper']

                    # write the scenario names and confidence interval titles
                    for ci in range(len(strings_to_write)):
                        row, column = reverse_inputs_if_required([1, s * 3 + 2 + ci], horizontal)
                        sheet.cell(row=row, column=column).value = strings_to_write[ci]

                    # write the columns of data
                    for y, year in enumerate(self.years_to_write):
                        for o in range(3):
                            row, column = reverse_inputs_if_required([y + 2, s * 3 + 2 + o], horizontal)
                            sheet.cell(row=row, column=column).value \
                                = self.uncertainty_centiles['epi'][scenario][inter][
                                o, t_k.find_first_list_element_at_least(self.interpolation_times_uncertainty, year)]

                # without uncertainty
                else:

                    # write scenario names across first row
                    row, column = reverse_inputs_if_required([1, s + 2], horizontal)
                    sheet.cell(row=row, column=column).value = t_k.capitalise_and_remove_underscore(scenario_name)

                    # write columns of data
                    for y, year in enumerate(self.years_to_write):
                        row, column = reverse_inputs_if_required([y + 2, s + 2], horizontal)
                        sheet.cell(row=row, column=column).value \
                            = self.model_runner.outputs['manual']['epi'][scenario][inter][
                                t_k.find_first_list_element_at_least(self.model_runner.outputs['manual']['epi'][
                                                                               scenario]['times'], year)]
            workbook.save(path)

        # code probably could bug because interventions can differ by scenario
        for inter in self.inputs.interventions_to_cost[0]:
            for cost_type in ['raw_cost_', 'inflated_cost_', 'discounted_cost_', 'discounted_inflated_cost_']:

                # make filename
                path = os.path.join(self.out_dir_project, cost_type + inter) + '.xlsx'

                # get active sheet
                workbook = xl.Workbook()
                sheet = workbook.active
                sheet.title = inter

                # write the year text cell
                sheet.cell(row=1, column=1).value = 'Year'

                # write the year text column
                for y, year in enumerate(self.years_to_write):
                    row, column = reverse_inputs_if_required([y + 2, 1], horizontal)
                    sheet.cell(row=row, column=column).value = year

                # cycle over scenarios
                for s, scenario in enumerate(scenarios):
                    scenario_name = t_k.find_scenario_string_from_number(scenario)

                    # scenario names
                    row, column = reverse_inputs_if_required([1, s + 2], horizontal)
                    sheet.cell(row=row, column=column).value = t_k.capitalise_and_remove_underscore(scenario_name)

                    # data columns
                    for y, year in enumerate(self.years_to_write):
                        row, column = reverse_inputs_if_required([y + 2, s + 2], horizontal)
                        sheet.cell(row=row, column=column).value \
                            = self.outputs['manual']['cost'][scenario][cost_type + inter][
                                t_k.find_first_list_element_at_least(self.outputs['manual']['cost'][scenario][
                                                                               'times'], year)]
                workbook.save(path)

    def write_docs_by_scenario(self):
        """
        Write word documents using the docx package. Writes with or without uncertainty according to whether Run
        uncertainty selected in the GUI. Currently only working for epidemiological outputs.
        """

        scenarios = [15] if self.run_mode == 'int_uncertainty' else self.scenarios
        for scenario in scenarios:

            # initialise document and table
            scenario_name = t_k.find_scenario_string_from_number(scenario)
            path = os.path.join(self.out_dir_project, scenario_name) + ".docx"
            document = Document()
            table = document.add_table(rows=len(self.years_to_write) + 1,
                                       cols=len(self.model_runner.epi_outputs_to_analyse) + 1)

            # for each epidemiological indicator
            for o, output in enumerate(self.model_runner.epi_outputs_to_analyse):

                # titles across the top
                row_cells = table.rows[0].cells
                row_cells[0].text = 'Year'
                row_cells[o + 1].text = t_k.capitalise_and_remove_underscore(output)

                # data columns
                for y, year in enumerate(self.years_to_write):

                    # write year column
                    row_cells = table.rows[y + 1].cells
                    row_cells[0].text = str(year)

                    # with uncertainty
                    if 'uncertainty' in self.run_mode:
                        point_lower_upper \
                            = tuple(self.uncertainty_centiles['epi'][scenario][output][
                                    0:3, t_k.find_first_list_element_at_least(self.interpolation_times_uncertainty,
                                                                              year)])
                        row_cells[o + 1].text = '%.1f\n(%.1f to %.1f)' % point_lower_upper

                    # without
                    else:
                        point = self.model_runner.outputs['manual']['epi'][scenario][output][
                            t_k.find_first_list_element_at_least(self.interpolation_times_uncertainty, year)]
                        row_cells[o + 1].text = '%.1f' % point
            document.save(path)

    def write_docs_by_output(self):
        """
        Write word documents using the docx package. Writes with or without uncertainty according to whether Run
        uncertainty selected in the GUI.
        """

        # write a new file for each output
        scenarios = [15] if self.run_mode == 'int_uncertainty' else self.scenarios
        for output in self.model_runner.epi_outputs_to_analyse:

            # initialise document, years of interest and table
            path = os.path.join(self.out_dir_project, output) + ".docx"
            document = Document()
            table = document.add_table(rows=len(self.years_to_write) + 1, cols=len(self.scenarios) + 1)

            for s, scenario in enumerate(scenarios):
                scenario_name = t_k.find_scenario_string_from_number(scenario)

                # outputs across the top
                row_cells = table.rows[0].cells
                row_cells[0].text = 'Year'
                row_cells[s + 1].text = t_k.capitalise_and_remove_underscore(scenario_name)

                for y, year in enumerate(self.years_to_write):
                    row_cells = table.rows[y + 1].cells
                    row_cells[0].text = str(year)

                    # with uncertainty
                    if 'uncertainty' in self.run_mode:
                        point_lower_upper \
                            = tuple(self.uncertainty_centiles['epi'][scenario][output][0:3,
                                    t_k.find_first_list_element_at_least(self.model_runner.outputs['manual'][
                                                                                   'epi'][scenario]['times'], year)])
                        row_cells[s + 1].text = '%.1f\n(%.1f to %.1f)' % point_lower_upper

                    # without
                    else:
                        point = self.model_runner.outputs['manual']['epi'][scenario][output][
                            t_k.find_first_list_element_at_least(self.model_runner.outputs['manual'][
                                                                           'epi'][scenario]['times'], year)]
                        row_cells[s + 1].text = '%.1f' % point
            document.save(path)

    def print_int_uncertainty_relative_change(self, year):
        """
        Print some text giving percentage change in output indicators relative to baseline under intervention
        uncertainty.

        Args:
            year: Year for the comparisons to be made against
        """

        scenario, changes = 15, {}
        for output in self.model_runner.epi_outputs_to_analyse:
            absolute_values \
                = self.uncertainty_centiles['epi'][scenario][output][0:3,
                  t_k.find_first_list_element_at_least(self.interpolation_times_uncertainty, year)]
            baseline = self.model_runner.outputs['manual']['epi'][0][output][
               t_k.find_first_list_element_at_least(self.model_runner.outputs['manual']['epi'][0]['times'], year)]
            changes[output] = [(i / baseline - 1.) * 1e2 for i in absolute_values]
            print(output + '\n%.1f\n(%.1f to %.1f)' % tuple(changes[output]))

    def print_average_costs(self):
        """
        Incompletely developed method to display the mean cost of an intervention over the course of its implementation
        under baseline conditions.
        """

        for scenario in self.scenarios:
            print('\n' + t_k.find_scenario_string_from_number(scenario))
            mean_cost = {}
            for inter in self.model_runner.interventions_to_cost[scenario]:
                print('\n' + inter)
                mean_cost[inter] \
                    = numpy.mean(self.model_runner.outputs['manual']['cost'][scenario]['raw_cost_' + inter])
                print('%.1f' % mean_cost[inter])
            print('total: %.1f' % sum(mean_cost.values()))

    ''' plotting methods '''

    def run_plotting(self):
        """
        Master plotting method to call all the methods that produce specific plots.
        """

        # plot mixing matrix whenever relevant
        if self.inputs.is_vary_force_infection_by_riskgroup and len(self.inputs.riskgroups) > 1:
            self.plot_mixing_matrix()

        # plot epidemiological outputs, overall, MDR-TB and risk groups
        if self.gui_inputs['output_epi_plots']:
            purposes = ['scenario', 'ci', 'progress', 'shaded'] if '_uncertainty' in self.run_mode else ['scenario']
            for purpose in purposes:
                self.plot_epi_outputs(self.gtb_available_outputs, purpose, 'main')
            if self.inputs.n_strains > 1:
                mdr_indicators = [ind + '_mdr' for ind in self.gtb_available_outputs if ind != 'notifications']
                mdr_indicators.append('perc_incidence_mdr')
                self.plot_epi_outputs(mdr_indicators, 'scenario', 'mdr-tb-related')
        if self.gui_inputs['output_by_subgroups']:
            for strata_type in ['agegroups', 'riskgroups']:
                outputs_to_plot, list_of_strata = ['incidence', 'mortality'], getattr(self.inputs, strata_type)
                if len(list_of_strata) > 1:
                    self.plot_epi_outputs(
                        [''.join(panel) for panel in itertools.product(outputs_to_plot, list_of_strata)],
                        'scenario', 'by_' + strata_type, grid=[len(outputs_to_plot), len(list_of_strata)], sharey='row')
            for strata_type in ['agegroups', 'riskgroups']:
                for fraction in [True, False]:
                    self.plot_stacked_epi_outputs('notifications', category_to_loop=strata_type, fraction=fraction)

        # plot scale-up functions
        if self.gui_inputs['output_scaleups']:
            self.plot_scaleup_vars()

        # plot economic outputs
        if self.gui_inputs['output_plot_economics']:
            self.plot_cost_coverage_curves()
            self.plot_cost_over_time()

        # plot compartment population sizes
        if self.gui_inputs['output_compartment_populations']:
            for category in ['agegroups', 'riskgroups', 'compartment_types']:
                for fraction in [True, False]:
                    self.plot_populations(category_to_loop=category, fraction=fraction)

        # make a flow-diagram
        if self.gui_inputs['output_flow_diagram']:
            self.model_runner.models[0].make_flow_diagram(
                os.path.join(self.out_dir_project, self.country + '_flow_diagram' + '.png'))

        # uncertainty figures
        if self.run_mode == 'epi_uncertainty' and self.gui_inputs['output_param_plots']:
            self.plot_param_histograms()
            self.plot_param_progression()
            self.plot_priors()

        # plot likelihood estimates
        if self.run_mode == 'epi_uncertainty' and self.gui_inputs['output_likelihood_plot']:
            self.plot_likelihoods()

    ''' epi outputs plotting '''

    def plot_epi_outputs(self, outputs, purpose, descriptor, grid=None, sharey='none'):
        """
        Produces the plot for the main outputs, loops over multiple scenarios.

        Args:
            outputs: A list of the outputs to be plotted
            purpose: Reason for plotting or type of plot, can be either 'scenario', 'ci_plot' or 'progress'
            descriptor: String for the filename and title of the plot
            grid: Shape of grid panels requested at call to method
            sharey: Whether to share the y-axis across rows of plots
        """

        # prelims
        fig, axes, max_dims, n_rows, n_cols \
            = initialise_figures_axes(len(outputs), requested_grid=grid, share_yaxis=sharey)
        start_time = self.inputs.model_constants['before_intervention_time'] \
            if self.run_mode == 'int_uncertainty' or (len(self.scenarios) > 1 and purpose == 'scenario') \
            else self.gui_inputs['plot_option_start_time']
        start_index, max_data_values = 0, {}
        scenarios, uncertainty_scenario = ([0, 15], 15) if self.run_mode == 'int_uncertainty' \
            else (self.scenarios, 0)

        # loop through output indicators
        for out, output in enumerate(outputs):
            axis = find_panel_grid_indices(axes, out, n_rows, n_cols)
            max_data_values[output] = []

            # overlay GTB data
            if self.gui_inputs['plot_option_overlay_gtb'] and output in self.gtb_available_outputs:
                max_data_values[output].append(self.plot_gtb_data_to_axis(
                    axis, output, start_time, self.gtb_indices[output] if output in self.gtb_indices else '',
                    gtb_ci_plot='hatch' if purpose == 'shaded' else 'patch'))

            # plot with uncertainty confidence intervals (median, lower, upper)
            if purpose == 'ci':
                for scenario in scenarios:
                    if self.run_mode == 'int_uncertainty' and scenario == 0:
                        continue
                    start_index = self.find_start_time_index(start_time, scenario)
                    max_data_values[output].append(
                        max(self.uncertainty_centiles['epi'][scenario][output][2, :][start_index:]))
                    for ci in range(3):
                        axis.plot(
                            self.interpolation_times_uncertainty[start_index:],
                            self.uncertainty_centiles['epi'][scenario][output][ci, :][start_index:],
                            color='k', label=None, linewidth=.7 if ci == 0 else .5, linestyle='-' if ci == 0 else '--')

            # plot progressive model run outputs for uncertainty analyses
            elif purpose == 'progress':
                runs_to_loop = self.inputs.n_samples if self.run_mode == 'int_uncertainty' \
                    else len(self.outputs['epi_uncertainty']['epi'][uncertainty_scenario][output])
                for run in range(runs_to_loop):
                    if run in self.accepted_indices or self.plot_rejected_runs or self.run_mode == 'int_uncertainty':
                        start_index = self.find_start_time_index(start_time, 0, by_run=True, run=run)
                        dotted = '.' if self.run_mode == 'epi_uncertainty' and run not in self.accepted_indices else '-'
                        colour = str(1. - float(run) / float(len(
                            self.outputs[self.run_mode]['epi'][uncertainty_scenario][output]))) \
                            if self.run_mode == 'epi_uncertainty' else '.4'
                        plot_data = self.outputs[self.run_mode]['epi'][uncertainty_scenario][output][run, start_index:]
                        max_data_values[output].append(max(plot_data))
                        axis.plot(self.outputs[self.run_mode]['epi'][uncertainty_scenario]['times'][run,
                                  start_index:], plot_data, color=colour, linestyle=dotted)

            # plot with shaded patches
            elif purpose == 'shaded':
                axis.patch.set_facecolor((1., 1., 1.))
                for side in ['top', 'bottom', 'left', 'right']:
                    axis.spines[side].set_color('.6')
                axis.grid(color='.8')
                start_index = self.find_start_time_index(start_time, 0)
                max_data_values[output].append(
                    max(self.uncertainty_centiles['epi'][uncertainty_scenario][output][-5, :][start_index:]))
                for i in range(self.model_runner.n_centiles_for_shading):
                    prop_progress = float(i) / float(self.model_runner.n_centiles_for_shading - 1)
                    patch_colour = (1. - prop_progress, 1. - prop_progress, 1 - prop_progress * .2)
                    patch = create_patch_from_list(
                        self.interpolation_times_uncertainty[start_index:],
                        self.uncertainty_centiles['epi'][uncertainty_scenario][output][i + 3, :][start_index:],
                        self.uncertainty_centiles['epi'][uncertainty_scenario][output][-i - 1, :][start_index:])
                    axis.add_patch(patches.Polygon(patch, color=patch_colour))

            # plot scenarios without uncertainty
            if purpose == 'scenario' or self.run_mode == 'int_uncertainty':
                scenarios_for_baseline = [0] if self.run_mode == 'int_uncertainty' else scenarios
                for scenario in scenarios_for_baseline:
                    start_index = self.find_start_time_index(start_time, scenario, purpose='scenario')
                    if self.run_mode == 'increment_comorbidity' and scenario != 0:
                        colour = (1. - (0.3 + self.inputs.comorbidity_prevalences[scenario]) * .2,
                                  1. - (0.3 + self.inputs.comorbidity_prevalences[scenario]),
                                  1. - (0.3 + self.inputs.comorbidity_prevalences[scenario]))
                        label = str(int(self.inputs.comorbidity_prevalences[scenario] * 1e2)) + '%'
                    else:
                        colour = self.colour_theme[scenario]
                        label = t_k.capitalise_and_remove_underscore(t_k.find_scenario_string_from_number(scenario))
                    # label = str(int(self.inputs.comorbidity_prevalences[scenario] * 1e2)) + '%' \
                    #     if self.run_mode == 'increment_comorbidity' \
                    #     else t_k.capitalise_and_remove_underscore(t_k.find_scenario_string_from_number(scenario))
                    max_data_values[output].append(max(self.outputs['manual']['epi'][scenario][output][start_index:]))
                    axis.plot(self.outputs['manual']['epi'][scenario]['times'][start_index:],
                              self.outputs['manual']['epi'][scenario][output][start_index:],
                              color=colour, linewidth=1.5, label=label,
                              zorder=1 if scenario else 4)

            # add plotting of End TB Targets
            if self.gui_inputs['plot_option_overlay_targets'] and (output == 'incidence' or output == 'mortality'):
                self.plot_targets_to_axis(axis, output)

            # finishing off axis and figure
            self.tidy_x_axis(axis, start_time, 2035., max_dims, labels_off=not last_row(out, n_rows, n_cols))
            self.tidy_y_axis(axis, output, max_dims, max_value=max(max_data_values[output]))
            axis.set_title(t_k.find_title_from_dictionary(output), fontsize=get_label_font_size(max_dims))
            if out == len(outputs) - 1 and purpose == 'scenario' and len(self.scenarios) > 1:
                add_legend_to_plot(axis, max_dims)
        self.finish_off_figure(fig, len(outputs), '_' + descriptor + '_epi_' + purpose,
                               'Epidemiological outputs'
                               + t_k.find_title_from_dictionary(descriptor, capital_first_letter=False)
                               + ', ' + t_k.capitalise_first_letter(self.country))

    def plot_targets_to_axis(self, axis, output, compare_gtb=False):
        """
        Plot End TB Target values to outputs axis.

        Args:
            axis: Plot axis to plot on to
            output: Output type
            compare_gtb: Whether to compare against reported (True) or modelled (False) outputs
        """

        uncertainty_scenario = 15 if self.run_mode == 'int_uncertainty' else 0

        # find the baseline value to compare the targets against
        if compare_gtb:
            base_value = self.inputs.original_data['notifications']['c_newinc'][2016] \
                if output == 'notifications' else self.inputs.original_data['gtb'][
                self.gtb_indices[output] + self.level_conversion_dict['point_estimate']][2016]
        elif self.run_mode == 'epi_uncertainty':
            base_value = self.uncertainty_centiles['epi'][0][output][0, t_k.find_first_list_element_at_least(
                self.interpolation_times_uncertainty, 2015.)]
        else:
            base_value = self.outputs['manual']['epi'][uncertainty_scenario][output][
                t_k.find_first_list_element_at_least(
                    self.outputs['manual']['epi'][uncertainty_scenario]['times'], 2015.)]

        # plot the milestones and targets
        plot_endtb_targets(axis, output, base_value, '.7')

    def plot_gtb_data_to_axis(self, ax, output, start_time, output_index, gtb_ci_plot='hatch'):
        """
        Method to plot the data loaded directly from the GTB report in the background.

        Args:
            ax: Axis for plotting
            output: String for output
            start_time: Starting time for plot axis
            output_index: String to index the GTB data
            gtb_ci_plot: How to display the confidence intervals of the GTB data
        Returns:
            The maximum value from the point estimate data being plotted
        """

        # prelims
        gtb_data, gtb_data_lists, gtb_index, colour, line_width, alpha = {}, {}, 0, '.2', 0., 1.

        # notifications
        if output == 'notifications':
            gtb_data['point_estimate'] = self.inputs.original_data['notifications']['c_newinc']
            gtb_data_lists.update(t_k.extract_dict_to_ordered_key_lists(gtb_data['point_estimate'], 'point_estimate'))
            gtb_index = t_k.find_first_list_element_at_least(gtb_data_lists['times'], start_time)

        # extract data other outputs
        else:
            for level in self.level_conversion_dict:
                gtb_data[level] = self.inputs.original_data['gtb'][output_index + self.level_conversion_dict[level]]
                gtb_data_lists.update(t_k.extract_dict_to_ordered_key_lists(gtb_data[level], level))
            gtb_index = t_k.find_first_list_element_at_least(gtb_data_lists['times'], start_time)

            # plot patch
            colour, hatch, fill, line_width, alpha = (self.gtb_patch_colours[output], None, True, 0., .5) \
                if gtb_ci_plot == 'patch' else ('.3', '/', False, .8, 1.)
            ax.add_patch(patches.Polygon(create_patch_from_list(gtb_data_lists['times'][gtb_index:],
                                                                gtb_data_lists['lower_limit'][gtb_index:],
                                                                gtb_data_lists['upper_limit'][gtb_index:]),
                                         color=colour, hatch=hatch, fill=fill, linewidth=0., alpha=alpha, zorder=4))

        # plot point estimates
        ax.plot(gtb_data['point_estimate'].keys()[gtb_index:], gtb_data['point_estimate'].values()[gtb_index:],
                color=colour, linewidth=.8, label=None, alpha=alpha)
        if gtb_ci_plot == 'hatch' and output != 'notifications':
            for limit in ['lower_limit', 'upper_limit']:
                ax.plot(gtb_data[limit].keys()[gtb_index:], gtb_data[limit].values()[gtb_index:],
                        color=colour, linewidth=line_width, label=None, alpha=alpha)

        return max(gtb_data['point_estimate'].values())

    def plot_populations(self, category_to_loop='agegroups', scenario=0, fraction=False, requirements=('',),
                         exclusions=('we all love futsal',)):
        """
        Plot population by the compartment or other category to which they belong. Doesn't necessarily work that well
        for compartments, where there are generally too many compartments for them to be easily visualised.

        Args:
            category_to_loop: Must be either 'compartment', 'agegroups' or 'riskgroups' to indicate category
            scenario: Generally 0 to indicate baseline scenario
            fraction: Boolean for whether we want population totals or fractions
            requirements: List of strings that must be present for compartment to be included in analysis
            exclusions: List of strings that indicate that compartment should be excluded from analysis
        """

        # prelims
        fig, ax, max_dim, n_rows, n_cols = initialise_figures_axes(1, room_for_legend=True)
        start_time = self.inputs.model_constants['plot_start_time']
        start_time_index = self.find_start_time_index(start_time, scenario)
        times = self.model_runner.outputs['manual']['epi'][scenario]['times'][start_time_index:]

        # get data
        cumulative_data = [0.] * len(times)
        current_data = self.sum_compartments_by_category(category_to_loop, scenario, start_time_index,
                                                         requirements=requirements, exclusions=exclusions)

        # plot patches and proxy by category
        for l, label in enumerate(current_data):
            previous_data, cumulative_data = increment_list_for_patch(current_data[label], cumulative_data)
            colour = self.colour_theme[l + 1]
            ax.fill_between(times, previous_data, cumulative_data, facecolor=colour, edgecolor=colour, alpha=.8)
            ax.plot([-1e2], [0.], color=colour, label=t_k.find_title_from_dictionary(label), linewidth=5.)  # proxy

        # finish off
        self.tidy_x_axis(ax, start_time, 2035., max_dim)
        self.tidy_y_axis(ax, '', max_dim, max_value=max(cumulative_data))
        ax.legend(bbox_to_anchor=(1.3, 1))
        filename = '_' + ('fraction' if fraction else 'population') + '_' + category_to_loop
        title = ('Fraction' if fraction else 'Size') + ' of population by ' \
            + t_k.find_title_from_dictionary(category_to_loop, capital_first_letter=False)
        for requirement in requirements:
            if requirement != '':
                filename += '_only' + requirement
                title += ',' + requirement + ' only'
        for exclusion in exclusions:
            if exclusion != 'we all love futsal':
                filename += '_exclude' + exclusion
                title += ', except ' + exclusion
        self.finish_off_figure(fig, 1, filename, title)

    def plot_stacked_epi_outputs(self, output='notifications', category_to_loop='agegroups', scenario=0,
                                 fraction=True):
        """
        Method to plot the proportion of notifications that come from various groups of the model. Particularly intended
        to keep an eye on the proportion of notifications occurring in the paediatric population (which WHO sometimes
        say should be around 15% in well-functioning TB programs).

        Args:
            output: Epidemiological output of interest
            category_to_loop: String of the model attribute of interest - can set to 'riskgroups'
            scenario: Scenario to take the outputs from (generally 0 for baseline)
            fraction: Boolean for whether to plot values as
        """

        # prelims
        fig, ax, max_dim, n_rows, n_cols = initialise_figures_axes(1, room_for_legend=True)
        strata = getattr(self.inputs, category_to_loop)
        start_time = self.inputs.model_constants['plot_start_time']
        start_time_index = self.find_start_time_index(start_time, scenario)
        times = self.model_runner.models[0].times[start_time_index:]
        cumulative_data = [0.] * len(times)

        for s, stratum in enumerate(strata):

            # get data
            current_data = self.outputs['manual']['epi'][0][output + stratum][start_time_index:]
            if fraction:
                current_data = t_k.elementwise_list_division(
                    current_data, self.outputs['manual']['epi'][0][output][start_time_index:])
            previous_data, cumulative_data = increment_list_for_patch(current_data, cumulative_data)

            # plot patch and proxy
            colour = self.colour_theme[s + 1]
            ax.fill_between(times, previous_data, cumulative_data, facecolors=colour, edgecolor=colour, alpha=.8)
            ax.plot([-1e2], [0.], color=colour, linewidth=5.,
                    label=t_k.turn_strat_into_label(stratum) if category_to_loop == 'agegroups'
                    else t_k.find_title_from_dictionary(stratum))  # proxy

        # finish off
        ax.legend(bbox_to_anchor=(1.3, 1))
        self.tidy_x_axis(ax, start_time, 2035., max_dim)
        self.tidy_y_axis(ax, 'prop_' if fraction else '', max_dim, max_value=max(cumulative_data))
        self.finish_off_figure(fig, 1, '_' + ('absolute_' if fraction else 'population_') + output + '_'
                               + category_to_loop,
                               ('Fraction of ' if fraction else 'Stacked absolute ') + output + ', by '
                               + t_k.find_title_from_dictionary(category_to_loop, capital_first_letter=False))

    ''' miscellaneous plotting method '''

    def plot_mixing_matrix(self):
        """
        Method to visualise the mixing matrix with bar charts.
        """

        # prelims
        fig, ax, max_dims, _, _ = initialise_figures_axes(1, room_for_legend=True)
        cumulative_data, bar_width, x_positions = list(numpy.zeros(len(self.inputs.riskgroups))), .7, []

        # plot bars
        for to, to_group in enumerate(self.inputs.riskgroups):
            previous_data, cumulative_data \
                = increment_list_for_patch(
                  [self.inputs.mixing[from_group][to_group] for from_group in self.inputs.riskgroups], cumulative_data)
            x_positions = numpy.linspace(.5, .5 + len(cumulative_data) - 1., len(cumulative_data))
            ax.bar(x_positions, cumulative_data, width=bar_width, bottom=previous_data, color=self.colour_theme[to],
                   label=t_k.find_title_from_dictionary(to_group))

        # locally managing x-axis, as plot type is a special case
        ax.set_xlim(.2, max(x_positions) + 1.)
        ax.set_xticks([x + bar_width / 2. for x in x_positions])
        ax.tick_params(axis='x', length=0.)
        ax.set_xticklabels([t_k.find_title_from_dictionary(group) for group in self.inputs.riskgroups],
                           fontsize=get_label_font_size(1))

        # finish off
        self.tidy_y_axis(ax, 'prop_', max_dims, max_value=1., space_at_top=0.)
        add_legend_to_plot(ax, max_dims)
        self.finish_off_figure(fig, max_dims, '_mixing', 'Source of contacts by risk group')

    ''' scale-up function plotting '''

    def plot_scaleup_vars(self):
        """
        Method that can be used to visualise each scale-up variable, not plotted against the data it is fit to and only
        on a single panel.
        """

        # prelims
        n_panels = 2 if self.gui_inputs['plot_option_vars_two_panels'] else 1
        vars_to_plot = self.model_runner.models[0].scaleup_fns.keys()
        if self.gui_inputs['plot_option_plot_all_vars']:
            vars_to_plot = t_k.combine_two_lists_no_duplicate(vars_to_plot, self.model_runner.models[0].vars)
        for var in vars_to_plot:
            fig, axes, max_dims, n_rows, n_cols = initialise_figures_axes(n_panels)
            for n_axis in range(n_panels):

                # find time to plot from and x-values
                start_time = float(self.gui_inputs['plot_option_start_time']) if n_axis == n_panels - 1 \
                    else self.inputs.model_constants['early_time']
                end_time = float(self.gui_inputs['plot_option_end_time']) if n_axis == n_panels - 1 \
                    else self.inputs.model_constants['scenario_end_time']

                # plot
                max_var = self.plot_scaleup_var_to_axis(axes[n_axis], [start_time, end_time], var)
                max_data = self.plot_scaleup_data_to_axis(axes[n_axis], [start_time, end_time], var)

                # clean up axes
                self.tidy_x_axis(axes[n_axis], start_time, end_time, max_dims)
                self.tidy_y_axis(axes[n_axis], var, max_dims, left_axis=n_axis % n_cols == 0,
                                 max_value=float(max([max_var, max_data])))

            self.finish_off_figure(fig, n_panels, '_' + var, var)

    def plot_scaleup_var_to_axis(self, axis, time_limits, var):
        """
        Add the scale-up var function output to an axis.

        Args:
            axis: The axis to add the line to
            time_limits: The limits of the horizontal axis in years
            var: String for the var to plot
        """

        maximum_values = []
        for scenario in reversed(self.scenarios):

            # if available as a scale-up function
            if var in self.model_runner.models[scenario].scaleup_fns:
                x_vals = numpy.linspace(time_limits[0], time_limits[1], int(1e3))
                y_vals = map(self.model_runner.models[scenario].scaleup_fns[var], x_vals)

            # otherwise if a different type of var, such as additional calculated ones
            else:
                start_time_index = self.find_start_time_index(time_limits[0], scenario)
                x_vals = self.model_runner.models[scenario].times[start_time_index:]
                y_vals = self.model_runner.models[scenario].get_var_soln(var)[start_time_index:]

            # plot and record the maximum value
            axis.plot(x_vals, y_vals, color=self.colour_theme[scenario],
                      label=t_k.capitalise_and_remove_underscore(t_k.find_scenario_string_from_number(scenario)))
            maximum_values.append(max(y_vals))
        return max(maximum_values)

    def plot_scaleup_data_to_axis(self, axis, time_limits, var):
        """
        Plot data that a scale-up function had been fitted to if it is in the desired range.

        Args:
            axis: Axis to plot onto
            time_limits: The limits of the horizontal axis in years
            var: String of the var to plot
        """

        if self.gui_inputs['plot_option_overlay_input_data'] and var in self.inputs.scaleup_data[0]:
            data_to_plot = {key: value for key, value in self.inputs.scaleup_data[0][var].items()
                            if int(time_limits[0]) <= key <= int(time_limits[1])}
            axis.scatter(data_to_plot.keys(), data_to_plot.values(), color=self.colour_theme[1], s=7, zorder=10)
            return max(data_to_plot.values()) if data_to_plot else 0.

    ''' economics plotting '''

    def plot_cost_coverage_curves(self):
        """
        Plots cost-coverage curves at times specified in the report times inputs in control panel.
        """

        # plot figures by scenario
        for scenario in self.scenarios:
            fig, axes, max_dim, n_rows, n_cols \
                = initialise_figures_axes(len(self.interventions_to_cost[scenario]), share_yaxis='row')
            fig.tight_layout()

            # subplots by program
            for p, program in enumerate(self.interventions_to_cost[scenario]):
                axis = find_panel_grid_indices(axes, p, n_rows, n_cols)
                end_value = 0.

                # generate times to plot cost-coverage curves at, inclusively (by adding small value to end time)
                times = numpy.arange(self.inputs.model_constants['cost_curve_start_time'],
                                     self.inputs.model_constants['cost_curve_end_time'] + .01,
                                     self.inputs.model_constants['cost_curve_step_time'])

                # plot costs versus coverage
                for t, time in enumerate(times):
                    coverage = numpy.arange(0., self.inputs.model_constants['econ_saturation_' + program], .02)
                    costs = [1e-3 * economics.get_cost_from_coverage(
                        cov, self.inputs.model_constants['econ_inflectioncost_' + program],
                        self.inputs.model_constants['econ_saturation_' + program],
                        self.inputs.model_constants['econ_unitcost_' + program],
                        self.model_runner.models[scenario].var_array[
                            t_k.find_first_list_element_at_least(self.model_runner.models[scenario].times, time),
                            self.model_runner.models[scenario].var_labels.index('popsize_' + program)])
                        for cov in coverage]
                    axis.plot(costs, coverage, label=str(int(time)),
                              color=(1. - float(t) / float(len(times)), 1. - float(t) / float(len(times)),
                                     1. - float(t) / float(len(times)) * .5))
                    end_value = max([end_value, max(costs)])

                # finish off axis
                axis.set_title(t_k.find_title_from_dictionary('program_prop_' + program),
                               fontsize=get_label_font_size(max_dim))
                self.tidy_x_axis(axis, 0., end_value, max_dim,
                                 x_label='Thousand $US' if last_row(p, n_rows, n_cols) else None)
                self.tidy_y_axis(axis, 'prop_', max_dim, max_value=1., left_axis=p % n_cols == 0, y_label='',
                                 y_lims=[0., 1.])
                if p == len(self.interventions_to_cost[scenario]) - 1:
                    add_legend_to_plot(axis, max_dim, location=4)

            self.finish_off_figure(fig, len(self.interventions_to_cost[scenario]),
                                   '_cost_coverage_' + t_k.find_scenario_string_from_number(scenario),
                                   'Cost coverage curves, ' + t_k.find_scenario_string_from_number(scenario))

    def plot_cost_over_time(self):
        """
        Method that produces plots for individual and cumulative program costs for each scenario as separate figures.
        Panels of figures are the different sorts of costs (i.e. whether discounting and inflation have been applied).
        """

        # separate figures for each scenario
        for scenario in self.scenarios:
            figs, axes, ax_dict, plot_types, n_rows, n_cols, max_dim \
                = {}, {}, {}, ['individual', 'stacked', 'relative'], 0, 0, 0
            for plot_type in plot_types:
                figs[plot_type], axes[plot_type], max_dim, n_rows, n_cols \
                    = initialise_figures_axes(len(self.model_runner.cost_types), share_yaxis='all')
            cost_times = self.outputs['manual']['cost'][scenario]['times']

            for c, cost_type in enumerate(self.model_runner.cost_types):
                for plot_type in plot_types:
                    ax_dict[plot_type] = find_panel_grid_indices(axes[plot_type], c, n_rows, n_cols)
                cumulative_data = [0.] * len(cost_times)
                for inter, intervention in enumerate(self.inputs.interventions_to_cost[scenario]):

                    # process data
                    current_data = self.outputs['manual']['cost'][scenario][cost_type + '_cost_' + intervention]
                    previous_data, cumulative_data = increment_list_for_patch(current_data, cumulative_data)
                    relative_data \
                        = [(d - current_data[t_k.find_first_list_element_above(
                            cost_times, self.inputs.model_constants['reference_time'])]) for d in current_data]

                    # plot lines and areas
                    colour = self.colour_theme[inter + 1]
                    ax_dict['individual'].plot(
                        cost_times, current_data, color=colour, label=t_k.find_title_from_dictionary(intervention))
                    ax_dict['relative'].plot(
                        cost_times, relative_data, color=colour, label=t_k.find_title_from_dictionary(intervention))
                    ax_dict['stacked'].fill_between(
                        cost_times, previous_data, cumulative_data, color=colour, edgecolor=colour, linewidth=.1)

                # finish off each axis
                for plot_type in plot_types:
                    ax_dict[plot_type].set_title(t_k.find_title_from_dictionary(cost_type),
                                                 fontsize=get_label_font_size(max_dim))
                    self.tidy_x_axis(ax_dict[plot_type], cost_times[0], cost_times[-1], max_dim)
                    self.tidy_y_axis(ax_dict[plot_type], 'cost', max_dim, left_axis=c % n_cols == 0, y_label='$US',
                                     allow_negative=True if plot_type == 'relative' else False)
                    if cost_type == self.model_runner.cost_types[-1]:
                        add_legend_to_plot(ax_dict[plot_type], max_dim)

            # finish off figure
            for plot_type in plot_types:
                self.finish_off_figure(
                    figs[plot_type], len(plot_types),
                    '_' + t_k.find_scenario_string_from_number(scenario) + '_' + plot_type,
                    t_k.find_scenario_string_from_number(scenario) + ', '
                    + t_k.find_title_from_dictionary(plot_type, capital_first_letter=False) + ' cost')

    ''' uncertainty plotting methods '''

    def plot_param_histograms(self):
        """
        Simple function to plot histograms of parameter values used in uncertainty analysis.
        """

        # prelims
        fig, axes, max_dims, n_rows, n_cols \
            = initialise_figures_axes(len(self.model_runner.outputs['epi_uncertainty']['all_parameters']))

        for p, param in enumerate(self.model_runner.outputs['epi_uncertainty']['all_parameters']):
            ax = find_panel_grid_indices(axes, p, n_rows, n_cols)

            # find data
            param_values = [self.model_runner.outputs['epi_uncertainty']['all_parameters'][param][i]
                            for i in self.accepted_no_burn_in_indices]

            # update param_values to account for the weights or the accepted runs
            param_values = t_k.apply_weighting(param_values, self.accepted_run_weights)

            # plot histogram for each parameter
            y, _, _ = ax.hist(param_values, bins=20, edgecolor='k')
            param_range_index = [i for i in range(len(self.inputs.param_ranges_unc))
                                 if self.inputs.param_ranges_unc[i]['key'] == param][0]
            param_range = self.inputs.param_ranges_unc[param_range_index]['bounds'][1] \
                - self.inputs.param_ranges_unc[param_range_index]['bounds'][0]
            ax.set_title(t_k.find_title_from_dictionary(param), fontsize=get_label_font_size(max_dims))

            # indicate parameter range
            for i in range(2):
                ax.axvline(x=self.inputs.param_ranges_unc[param_range_index]['bounds'][i], color='.3', linestyle=':')

            # finish off axes and figure
            self.tidy_x_axis(ax, self.inputs.param_ranges_unc[param_range_index]['bounds'][0] - param_range * .1,
                             self.inputs.param_ranges_unc[param_range_index]['bounds'][1] + param_range * .1, max_dims)
            self.tidy_y_axis(ax, '', max_dims, max_value=max(y))
        self.finish_off_figure(fig, len(self.model_runner.outputs['epi_uncertainty']['all_parameters']),
                               '_parameter_hist', 'Parameter histograms')

    def plot_param_progression(self):
        """
        Plot accepted parameter progress over time against run sequence.
        """

        n_params = len(self.model_runner.outputs['epi_uncertainty']['all_parameters'])
        fig, axes, max_dims, n_rows, n_cols = initialise_figures_axes(n_params)
        for p, param in enumerate(self.model_runner.outputs['epi_uncertainty']['all_parameters']):
            ax = find_panel_grid_indices(axes, p, n_rows, n_cols)
            data = [self.model_runner.outputs['epi_uncertainty']['all_parameters'][param][i]
                    for i in self.accepted_no_burn_in_indices]
            ax.plot(range(1, len(data) + 1), data)
            param_range_index = [i for i in range(len(self.inputs.param_ranges_unc))
                                 if self.inputs.param_ranges_unc[i]['key'] == param][0]
            ax.set_title(t_k.find_title_from_dictionary(param), fontsize=get_label_font_size(max_dims))
            for i in range(2):
                ax.plot([1, len(data) + 1], [self.inputs.param_ranges_unc[param_range_index]['bounds'][i]] * 2,
                        color='.3', linestyle=':')
            self.tidy_x_axis(ax, 1, n_params + 1, max_dims, labels_off=not last_row(p, n_rows, n_cols))
            param_range = self.inputs.param_ranges_unc[param_range_index]['bounds'][1] \
                - self.inputs.param_ranges_unc[param_range_index]['bounds'][0]
            self.tidy_y_axis(ax, '', max_dims, max_value=max(data),
                             y_lims=(self.inputs.param_ranges_unc[param_range_index]['bounds'][0] - param_range * .1,
                                     self.inputs.param_ranges_unc[param_range_index]['bounds'][1] + param_range * .1))
        self.finish_off_figure(fig, n_params, '_parameter_series', 'Parameter progression')

    def plot_priors(self):
        """
        Function to plot the prior distributions that are logged to get the prior contribution to the acceptance
        probability in the epidemiological uncertainty running.
        """

        fig, axes, max_dims, n_rows, n_cols = initialise_figures_axes(len(self.model_runner.inputs.param_ranges_unc))
        n_plot_points, x_values, y_values, description = 1000, [], [], None
        for p, param in enumerate(self.model_runner.inputs.param_ranges_unc):

            # find values to plot
            distribution, lower, upper = param['distribution'], param['bounds'][0], param['bounds'][1]
            if distribution == 'uniform':
                x_values, y_values = [lower, upper], [1. / (upper - lower)] * 2
                description = t_k.capitalise_first_letter(distribution)
            elif distribution == 'beta_2_2':
                x_values = numpy.linspace(0., 1., n_plot_points)
                y_values = [stats.beta.pdf(x, 2., 2.) for x in x_values]
                description = t_k.find_title_from_dictionary(distribution)
            elif distribution == 'beta_mean_stdev':
                x_values = numpy.linspace(0., 1., n_plot_points)
                alpha_value = ((1. - param['additional_params'][0]) / param['additional_params'][1] ** 2. - 1.
                               / param['additional_params'][0]) * param['additional_params'][0] ** 2.
                beta_value = alpha_value * (1. / param['additional_params'][0] - 1.)
                y_values = [stats.beta.pdf(x, alpha_value, beta_value) for x in x_values]
                description = 'Beta, params:\n%.2g, %.2g' % (alpha_value, beta_value)
            elif distribution == 'beta_params':
                x_values = numpy.linspace(0., 1., n_plot_points)
                y_values = [stats.beta.pdf(x, param['additional_params'][0], param['additional_params'][1])
                            for x in x_values]
                description \
                    = 'Beta, params:\n%.2g, %.2g' % (param['additional_params'][0], param['additional_params'][1])
            elif distribution == 'gamma_mean_stdev':
                x_values = numpy.linspace(lower, upper, n_plot_points)
                alpha_value = (param['additional_params'][0] / param['additional_params'][1]) ** 2.
                beta_value = param['additional_params'][1] ** 2. / param['additional_params'][0]
                y_values = [stats.gamma.pdf(x, alpha_value, scale=beta_value) for x in x_values]
                description = 'Gamma, params:\n%.2g, %.2g' % (alpha_value, beta_value)
            elif distribution == 'gamma_params':
                x_values = numpy.linspace(lower, upper, n_plot_points)
                y_values = [stats.gamma.pdf(x, param['additional_params'][0]) for x in x_values]
                description = 'Gamma, params:\n%.2g' % param['additional_params'][0]

            # plot
            ax = find_panel_grid_indices(axes, p, n_rows, n_cols)
            ax.plot(x_values, y_values)

            # finish off axes and figure
            ax.set_title(t_k.find_title_from_dictionary(param['key']), fontsize=get_label_font_size(max_dims))
            ax.text(lower + .05, max(y_values) / 2., description, fontsize=get_label_font_size(max_dims))
            self.tidy_x_axis(ax, x_values[0], x_values[-1], max_dims)
            self.tidy_y_axis(ax, '', max_dims, max_value=max(y_values))
            ax.set_ylim(bottom=0.)
        self.finish_off_figure(fig, len(self.model_runner.inputs.param_ranges_unc), '_priors',
                               'Parameter prior distributions')

    def plot_likelihoods(self):
        """
        Method to plot likelihoods over runs, differentiating accepted and rejected runs to illustrate progression.
        """

        fig, ax, max_dims, n_rows, n_cols = initialise_figures_axes(1)

        # plot rejected values from the previous acceptance to current rejection
        for i in self.outputs['epi_uncertainty']['rejected_indices']:
            last_acceptance_before = [j for j in self.accepted_indices if j < i][-1]
            ax.plot([last_acceptance_before, i],
                    [self.outputs['epi_uncertainty']['loglikelihoods'][last_acceptance_before],
                     self.outputs['epi_uncertainty']['loglikelihoods'][i]], marker='o', linestyle='--',
                    color=self.colour_theme[2])

        # plot accepted values
        ax.plot(self.accepted_indices,
                [self.outputs['epi_uncertainty']['loglikelihoods'][i] for i in self.accepted_indices],
                marker='o', color=self.colour_theme[1])

        # finish off
        fig.suptitle('Progression of likelihood', fontsize=self.title_size)
        self.tidy_x_axis(ax, 0, len(self.accepted_indices), max_dims, x_label='Runs')
        self.tidy_y_axis(ax, 'likelihood', max_dims, y_label='Log likelihood')
        ax.set_ylabel('Log likelihood', fontsize=get_label_font_size(max_dims), labelpad=1)
        self.finish_off_figure(fig, 1, '_likelihoods', 'Likelihood progression')

    ''' miscellaneous '''

    def open_output_directory(self, out_dir_project):
        """
        Opens the directory into which all the outputs have been placed.
        """

        operating_system = platform.system()
        if 'Windows' in operating_system:
            os.system('start  ' + out_dir_project)
        elif 'Darwin' in operating_system:
            os.system('open  ' + out_dir_project)
